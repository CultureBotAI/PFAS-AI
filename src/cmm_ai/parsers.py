"""File parsers for converting various formats to text/TSV.

This module provides parsers for:
- Excel files (xlsx/xls) to TSV
- Word documents (docx) to text
- PDF files to text
"""

import csv
import io
import re
from pathlib import Path
from typing import Optional, Union, List

import pandas as pd
from docx import Document
import PyPDF2


def sanitize_filename(filename: str) -> str:
    """Replace spaces and other problematic characters with underscores.
    
    Args:
        filename: Original filename
        
    Returns:
        Sanitized filename with spaces replaced by underscores
        
    Examples:
        >>> sanitize_filename("My File Name.txt")
        'My_File_Name.txt'
        >>> sanitize_filename("Data with   spaces.tsv")
        'Data_with_spaces.tsv'
    """
    # Replace spaces and consecutive whitespace with single underscores
    sanitized = re.sub(r'\s+', '_', filename)
    return sanitized


def xlsx_to_tsv(
    input_path: Union[str, Path], 
    output_path: Optional[Union[str, Path]] = None,
    sheet_name: Union[str, int, None] = 0
) -> str:
    """Convert Excel file to TSV format.
    
    Args:
        input_path: Path to input Excel file
        output_path: Path for output TSV file. If None, returns TSV string
        sheet_name: Sheet to convert (name, index, or None for all sheets)
        
    Returns:
        TSV content as string if output_path is None, otherwise empty string
        
    Examples:
        >>> # Example with mock file
        >>> content = xlsx_to_tsv("test.xlsx")  # doctest: +SKIP
        >>> print(content.split('\\n')[0])  # doctest: +SKIP
        column1	column2	column3
    """
    input_path = Path(input_path)
    
    # Read Excel file
    if sheet_name is None:
        # Read all sheets
        dfs = pd.read_excel(input_path, sheet_name=None)
        all_tsvs = []
        
        for name, df in dfs.items():
            tsv_buffer = io.StringIO()
            df.to_csv(tsv_buffer, sep='\t', index=False)
            all_tsvs.append(f"# Sheet: {name}\n{tsv_buffer.getvalue()}")
        
        tsv_content = "\n\n".join(all_tsvs)
    else:
        # Read single sheet
        df = pd.read_excel(input_path, sheet_name=sheet_name)
        tsv_buffer = io.StringIO()
        df.to_csv(tsv_buffer, sep='\t', index=False)
        tsv_content = tsv_buffer.getvalue()
    
    # Save or return
    if output_path:
        output_path = Path(output_path)
        # Sanitize the filename to replace spaces with underscores
        sanitized_name = sanitize_filename(output_path.name)
        output_path = output_path.parent / sanitized_name
        output_path.write_text(tsv_content)
        return ""
    else:
        return tsv_content


def docx_to_text(
    input_path: Union[str, Path], 
    output_path: Optional[Union[str, Path]] = None,
    preserve_paragraphs: bool = True
) -> str:
    """Convert Word document to plain text.
    
    Args:
        input_path: Path to input docx file
        output_path: Path for output text file. If None, returns text string
        preserve_paragraphs: Keep paragraph separations with double newlines
        
    Returns:
        Text content as string if output_path is None, otherwise empty string
        
    Examples:
        >>> # Example with mock file
        >>> text = docx_to_text("document.docx")  # doctest: +SKIP
        >>> print(text[:50])  # doctest: +SKIP
        This is the first paragraph of the document.
    """
    input_path = Path(input_path)
    
    # Read docx file
    doc = Document(input_path)
    
    # Extract text from paragraphs
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # Join paragraphs
    if preserve_paragraphs:
        text_content = "\n\n".join(paragraphs)
    else:
        text_content = "\n".join(paragraphs)
    
    # Extract text from tables
    tables_text = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append("\t".join(row_data))
        if table_data:
            tables_text.append("\n".join(table_data))
    
    if tables_text:
        text_content += "\n\n# Tables\n\n" + "\n\n".join(tables_text)
    
    # Save or return
    if output_path:
        output_path = Path(output_path)
        # Sanitize the filename to replace spaces with underscores
        sanitized_name = sanitize_filename(output_path.name)
        output_path = output_path.parent / sanitized_name
        output_path.write_text(text_content)
        return ""
    else:
        return text_content


def pdf_to_text(
    input_path: Union[str, Path], 
    output_path: Optional[Union[str, Path]] = None,
    page_numbers: Optional[List[int]] = None
) -> str:
    """Convert PDF to plain text.
    
    Args:
        input_path: Path to input PDF file
        output_path: Path for output text file. If None, returns text string
        page_numbers: List of page numbers to extract (0-indexed). None for all pages
        
    Returns:
        Text content as string if output_path is None, otherwise empty string
        
    Examples:
        >>> # Example with mock file
        >>> text = pdf_to_text("document.pdf", page_numbers=[0])  # doctest: +SKIP
        >>> print(text[:50])  # doctest: +SKIP
        This is the text from the first page of the PDF.
    """
    input_path = Path(input_path)
    
    # Read PDF file
    text_parts = []
    
    with open(input_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        # Determine pages to extract
        if page_numbers is None:
            pages_to_extract = range(len(pdf_reader.pages))
        else:
            pages_to_extract = page_numbers
        
        # Extract text from each page
        for page_num in pages_to_extract:
            if 0 <= page_num < len(pdf_reader.pages):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"# Page {page_num + 1}\n\n{text}")
    
    text_content = "\n\n".join(text_parts)
    
    # Save or return
    if output_path:
        output_path = Path(output_path)
        # Sanitize the filename to replace spaces with underscores
        sanitized_name = sanitize_filename(output_path.name)
        output_path = output_path.parent / sanitized_name
        output_path.write_text(text_content)
        return ""
    else:
        return text_content


def parse_file(
    input_path: Union[str, Path],
    output_path: Optional[Union[str, Path]] = None
) -> str:
    """Parse file based on extension and convert to text/TSV.
    
    Automatically detects file type and applies appropriate parser:
    - .xlsx, .xls -> TSV
    - .docx -> text
    - .pdf -> text
    
    Args:
        input_path: Path to input file
        output_path: Path for output file. Extension determines format.
                    If None, returns content as string
        
    Returns:
        Parsed content as string if output_path is None, otherwise empty string
        
    Raises:
        ValueError: If file extension is not supported
        
    Examples:
        >>> # Example with mock files
        >>> content = parse_file("data.xlsx")  # doctest: +SKIP
        >>> content = parse_file("report.pdf")  # doctest: +SKIP
    """
    input_path = Path(input_path)
    extension = input_path.suffix.lower()
    
    if extension in ['.xlsx', '.xls']:
        return xlsx_to_tsv(input_path, output_path)
    elif extension == '.docx':
        return docx_to_text(input_path, output_path)
    elif extension == '.pdf':
        return pdf_to_text(input_path, output_path)
    else:
        raise ValueError(f"Unsupported file extension: {extension}")